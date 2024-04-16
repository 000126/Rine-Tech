from django.shortcuts import render, redirect, HttpResponse
from django.contrib import messages
from . models import *
import pandas as pd
import docx2txt
import openpyxl
from .models import ExcelDocument
from django.http import JsonResponse


def index(request):
    return render(request, 'index.html')


def home(request):
    if not request.user.is_authenticated:
        messages.warning(request, "Please Login and Try Again")
        return redirect('/login')
    # Template = Template.objects.all()
    # if request.method == 'POST':
    #     temp_na = request.POST.get('template_name')
    #     temp_type = request.POST.get('template_type')
    #     p_code = request.POST.get('P_code')
    #     batch_num = request.POST.get('batch_num')
    #     a_r = request.POST.get('a_r number')
    #     Template = Template.objects.filter(
    #         template=temp_na, template_type=temp_type, project_code=p_code, batch_no=batch_num, a_r_no=a_r)
    # context = {
    #     'Temp': Template
    # }
    return render(request, "home.html",)


def temp(request):
    return render(request, "awc.html")


# @csrf_exempt
# def get_data(request):
#     # Process the request and return data as JSON
#     data = {'message': 'Data from Django server'}
#     return JsonResponse(data)


# def new(request):
#     return render(request, 'add_new_temp.html')

def new_template(request):
    if request.method == "POST":
        template_name = request.POST.get('templatename')
        template_type = request.POST.get('category')
        template_number = request.POST.get('templatenumber')
        uploaded_file = request.FILES.get('file')
        if uploaded_file and uploaded_file.name.endswith('.xls') or uploaded_file.name.endswith('.xlsx'):
            data = f"{template_name}-{template_type}-{template_number}"
            saved = ExcelDocument(field=data, uploaded_files=uploaded_file)
            saved.save()
            messages.success(request, "Uploaded Successfully")
            return redirect('/add_new_template')
        else:
            messages.warning(request, "Please upload a valid Excel file.")
            return render(request, 'add_new_temp.html')
    return render(request, "add_new_temp.html")

# to upload file by admin only


def file_list(request):
    files = ExcelDocument.objects.all()
    return render(request, 'file_list.html', {'files': files})


# for excel file upload
def view_excel(request, excel_id):
    if not request.user.is_authenticated:
        messages.error(request, "Please Login and Try Again")
        return redirect('/auth/login/')
    excel_document = ExcelDocument.objects.get(pk=excel_id)

    # Read Excel data using pandas
    excel_data = pd.read_excel(excel_document.uploaded_files.path)

    # Convert the DataFrame to HTML
    excel_html = excel_data.to_html(
        classes='table table-bordered', index=True,)

    return render(request, 'view_excel.html', {'excel_document': excel_document, 'excel_html': excel_html})

#view word
from docx import Document

# def upload_file(request):
#     if request.method == 'POST' and request.FILES.get('file'):
#         uploaded_file = request.FILES['file']

#         # Save the uploaded file
#         file_instance = WordDocument(file=uploaded_file)
#         file_instance.save()

#         # Process the Word file (you may need to install python-docx)
#         doc = Document(uploaded_file)
#         content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

#         return render(request, 'render_word.html', {'content': content})

#     return render(request, 'upload_file.html')

def upload_word_file(request):
    if request.method == "POST":
        template_name = request.POST.get('templatename')
        template_type = request.POST.get('category')
        template_number = request.POST.get('templatenumber')
        uploaded_file = request.FILES.get('file')
        if uploaded_file and uploaded_file.name.endswith('.docx,.doc'): 
            data = f"{template_name}-{template_type}-{template_number}"
            saved = WordDocument(title=data, uploaded_files=uploaded_file)
            saved.save()
            messages.success(request, "Uploaded Successfully")
            return redirect('list_word_files')
        else:
            messages.warning(request, "Please upload a valid Ms Word file.")
            return render(request, 'add_new_temp.html')
    return render(request, "upload_word_file.html")

# def upload_word_file(request):
#     if request.method == 'POST' and 'word_file' in request.FILES:
#         title = request.POST.get('title', '')
#         word_file = request.FILES['word_file']
#         instance = WordDocument(title=title, word_file=word_file)
#         instance.save()
#         return redirect('list_word_files')

#     return render(request, 'upload_word_file.html')


def render_word_file(request, document_id):
    word_document = WordDocument.objects.get(pk=document_id)
    text_content = docx2txt.process(word_document.word_file.path)
    return render(request, 'render_word_file.html', {'text_content': text_content})


def list_word_files(request):
    word_documents = WordDocument.objects.all()
    return render(request, 'list_word_files.html', {'word_documents': word_documents})